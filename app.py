# -*- coding: utf-8 -*-
"""Vercel 版（Flask + 生成 docx + 打包 zip + 上传到 Vercel Blob）

你当前项目在本机运行没问题，但部署到 Vercel 需要注意两点：

1) Vercel Functions 的文件系统大多是只读的（仅 /tmp 可写），并且函数是无状态的。
   所以不能像本地一样把 zip 写到项目目录再通过 /download 路由二次下载。

2) Vercel Functions 的请求/响应体有 4.5MB 上限。
   生成的 zip 往往会超过 4.5MB，因此推荐把 zip 上传到 Vercel Blob，
   再把 Blob 的 URL 返回给前端下载。

本文件在 app_fixed.py 的基础上做了“Vercel 适配”：
- 输出目录改到临时目录（tempfile.gettempdir()）
- /api/generate 生成 zip 后上传到 Vercel Blob（需要配置 BLOB_READ_WRITE_TOKEN）
- 返回 data.zip_url 为完整的 https URL（Blob URL）

本地仍可运行：python app_vercel.py
部署到 Vercel：把本文件改名为 app.py（或按 Vercel Flask 文档指定入口文件）。
"""

from __future__ import annotations

import datetime
import os
import re
import shutil
import tempfile
import uuid
import zipfile
from typing import Dict, Iterable, List, Optional, Tuple

from flask import Flask, abort, jsonify, request
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm


# ============ 基础配置 ============
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
DOCX_TEMPLATE_ROOT = os.path.join(BASE_DIR, "templates")  # docx 模板根目录

# Vercel 函数文件系统多数只读，使用系统临时目录（Vercel 上为 /tmp）
OUTPUT_ROOT = os.path.join(tempfile.gettempdir(), "fund-docx-output")
os.makedirs(OUTPUT_ROOT, exist_ok=True)

# Flask 不使用 Jinja 的 HTML templates（避免与 docx 模板目录 templates/ 冲突）
app = Flask(__name__)

# 匹配 {{占位符}} 的正则（占位符名可含中文）
PH_RE = re.compile(r"{{(.*?)}}")


# ============ 可选：下拉框配置（按需自行改） ============
SELECT_OPTIONS = {
    "基金类型": [
        "股票型基金",
        "混合型基金",
        "债券型基金",
        "货币市场基金",
        "基金中基金（FOF）",
        "QDII",
        "商品型基金",
        "REITs",
        "其他",
    ],
    "上市交易所": [
        "上海证券交易所",
        "深圳证券交易所",
        "北京证券交易所",
        "香港交易所",
        "其他",
    ],
    "ETF类型": [
        "跨市场ETF",
        "单市场ETF",
        "跨境ETF",
        "债券ETF",
        "商品ETF",
        "其他",
    ],
    "指数公司名称": [
        "中证指数有限公司",
        "深圳证券信息有限公司",
        "恒生指数有限公司",
        "其他",
    ],
    "指数公司简称": [
        "中证指数",
        "深证信息",
        "恒生指数",
        "其他",
    ],
}

# 如果占位符名包含这些关键词，默认用 textarea（长文本）
LONG_TEXT_KEYWORDS = ["风险", "揭示", "策略", "分析", "简介", "介绍", "说明", "情况", "内容"]


# ============ 通用工具函数（与 app_fixed.py 基本一致，略） ============
def _safe_join_under_root(root: str, user_path: str) -> str:
    """防止路径穿越：只允许 root 目录下的子路径"""
    user_path = (user_path or "").strip()
    joined = os.path.abspath(os.path.join(root, user_path))
    root_abs = os.path.abspath(root)
    if not (joined == root_abs or joined.startswith(root_abs + os.sep)):
        raise ValueError("非法路径")
    return joined


def _sanitize_folder_name(name: str, max_len: int = 80) -> str:
    name = (name or "").strip()
    name = re.sub(r"[\\/:*?\"<>|]+", "_", name)
    name = name.strip(" .")
    name = re.sub(r"\s+", " ", name)
    if not name:
        name = "output"
    if len(name) > max_len:
        name = name[:max_len].rstrip(" .")
    return name


def _ensure_unique_subdir(parent: str, desired: str) -> Tuple[str, str]:
    desired = _sanitize_folder_name(desired)
    path = os.path.join(parent, desired)
    if not os.path.exists(path):
        os.makedirs(path, exist_ok=True)
        return path, desired

    counter = 2
    while True:
        name = f"{desired}({counter})"
        path2 = os.path.join(parent, name)
        if not os.path.exists(path2):
            os.makedirs(path2, exist_ok=True)
            return path2, name
        counter += 1


def list_product_types() -> List[str]:
    if not os.path.isdir(DOCX_TEMPLATE_ROOT):
        return []
    items = []
    for name in os.listdir(DOCX_TEMPLATE_ROOT):
        full = os.path.join(DOCX_TEMPLATE_ROOT, name)
        if os.path.isdir(full) and not name.startswith("."):
            items.append(name)
    items.sort()
    return items


def get_product_folder(product_type: str) -> str:
    if product_type in (None, "", "__root__"):
        return DOCX_TEMPLATE_ROOT
    return _safe_join_under_root(DOCX_TEMPLATE_ROOT, product_type)


def list_docx_files(folder: str) -> List[str]:
    if not os.path.isdir(folder):
        return []
    files = []
    for fn in os.listdir(folder):
        if fn.startswith("~$"):
            continue
        if fn.lower().endswith(".docx"):
            files.append(fn)
    files.sort()
    return files


def extract_placeholders_from_docx(docx_path: str) -> List[str]:
    found = set()
    try:
        doc = Document(docx_path)
    except Exception:
        return []

    def iter_paragraphs(d: Document):
        for p in d.paragraphs:
            yield p
        for t in d.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
        for s in d.sections:
            for hf in [s.header, s.footer]:
                for p in hf.paragraphs:
                    yield p
                for t in hf.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                yield p

    for p in iter_paragraphs(doc):
        for m in PH_RE.finditer(p.text or ""):
            key = (m.group(1) or "").strip()
            if key:
                found.add(key)

    return sorted(found)


def build_schema(product_type: str) -> Dict:
    folder = get_product_folder(product_type)
    tpl_files = list_docx_files(folder)
    all_keys: List[str] = []
    key_set = set()
    for fn in tpl_files:
        keys = extract_placeholders_from_docx(os.path.join(folder, fn))
        for k in keys:
            if k not in key_set:
                key_set.add(k)
                all_keys.append(k)

    fields = []
    for k in all_keys:
        ftype = "text"
        if any(kw in k for kw in LONG_TEXT_KEYWORDS):
            ftype = "textarea"
        if k in SELECT_OPTIONS:
            ftype = "select"
        fields.append(
            {
                "key": k,
                "label": k,
                "type": ftype,
                "options": SELECT_OPTIONS.get(k, []),
                "required": False,
                "group": "默认",
                "hint": "",
            }
        )

    return {
        "product_type": product_type,
        "templates": tpl_files,
        "fields": fields,
        "meta": {"generated_by": "fund-docx"},
    }


def build_docx_mapping(
    all_keys: Iterable[str],
    values_raw: Dict,
    blank_unfilled: bool,
    image_keys: Optional[Iterable[str]] = None,
) -> Dict[str, Optional[str]]:
    image_keys = set(image_keys or [])
    mapping: Dict[str, Optional[str]] = {}
    for k in all_keys:
        if k in image_keys:
            mapping[k] = None
            continue
        if k in values_raw:
            v = values_raw.get(k)
            if v is None:
                mapping[k] = "" if blank_unfilled else None
            else:
                mapping[k] = str(v)
        else:
            mapping[k] = "" if blank_unfilled else None
    return mapping


def apply_placeholders_to_filename(filename: str, mapping: Dict[str, Optional[str]], blank_unfilled: bool) -> str:
    def repl(m: re.Match) -> str:
        k = (m.group(1) or "").strip()
        if k not in mapping:
            return "" if blank_unfilled else m.group(0)
        v = mapping[k]
        if v is None:
            return m.group(0)
        return v

    out = PH_RE.sub(repl, filename)
    out = re.sub(r"\s+", " ", out).strip()
    out = re.sub(r"[\\/:*?\"<>|]+", "_", out)
    return out or "output.docx"


# ---------- XML 级别替换（保留格式） ----------
def _iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for s in doc.sections:
        for hf in [s.header, s.footer]:
            for p in hf.paragraphs:
                yield p
            for t in hf.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p


def _get_run_text_nodes(paragraph) -> List[OxmlElement]:
    nodes: List[OxmlElement] = []
    for r in paragraph.runs:
        for t in r._r.iter(qn("w:t")):
            nodes.append(t)
    return nodes


def _paragraph_text_with_positions(paragraph) -> Tuple[str, List[Tuple[OxmlElement, int, int]]]:
    nodes = _get_run_text_nodes(paragraph)
    full = []
    positions = []
    cur = 0
    for n in nodes:
        txt = n.text or ""
        full.append(txt)
        start = cur
        end = cur + len(txt)
        positions.append((n, start, end))
        cur = end
    return "".join(full), positions


def _insert_line_break_after(text_node: OxmlElement):
    br = OxmlElement("w:br")
    text_node.addnext(br)


def replace_placeholders_in_doc(doc: Document, mapping: Dict[str, Optional[str]]):
    for p in _iter_all_paragraphs(doc):
        full_text, positions = _paragraph_text_with_positions(p)
        if "{{" not in full_text:
            continue

        matches = list(PH_RE.finditer(full_text))
        if not matches:
            continue

        # 从后往前替换，避免索引偏移
        for m in reversed(matches):
            key = (m.group(1) or "").strip()
            if key not in mapping:
                continue
            rep = mapping[key]
            if rep is None:
                continue

            start, end = m.start(), m.end()
            # 找到覆盖该占位符的节点范围
            left_idx = right_idx = None
            for i, (_, s, e) in enumerate(positions):
                if left_idx is None and s <= start < e:
                    left_idx = i
                if s < end <= e:
                    right_idx = i
                    break

            if left_idx is None or right_idx is None:
                continue

            left_node, ls, le = positions[left_idx]
            right_node, rs, re = positions[right_idx]

            left_text = left_node.text or ""
            right_text = right_node.text or ""

            prefix = left_text[: start - ls]
            suffix = right_text[end - rs :]

            # 清空中间节点
            for j in range(left_idx + 1, right_idx):
                positions[j][0].text = ""

            # 把替换内容写入 left_node，保持其原 run 样式
            # 多行用 <w:br/> 表示，避免破坏段落/列表样式
            rep_lines = (rep or "").split("\n")
            left_node.text = prefix + (rep_lines[0] if rep_lines else "")
            cur_node = left_node
            for line in rep_lines[1:]:
                _insert_line_break_after(cur_node)
                # 新建一个 w:t 节点，复用同一个 run（最大化保留样式）
                new_t = OxmlElement("w:t")
                new_t.text = line
                cur_node.addnext(new_t)
                cur_node = new_t

            # right_node 只保留 suffix
            if right_idx == left_idx:
                # 同一个节点：已写入 prefix+rep，继续拼 suffix
                cur_node.text = (cur_node.text or "") + suffix
            else:
                right_node.text = suffix


# ---------- 图片替换（可选：与 app_fixed.py 同逻辑，只保留最小实现） ----------
# ---------- 图片替换（与 app_fixed.py 同逻辑） ----------
def _resolve_local_path(base_dir: str, p: Optional[str]) -> Optional[str]:
    if not p:
        return None
    p = str(p).strip().strip('"').strip("'")
    if not p:
        return None

    # 支持 static/... 或 presets/... 这类相对路径
    if not os.path.isabs(p):
        return os.path.abspath(os.path.join(base_dir, p))
    return p


def _iter_all_paragraphs(doc: Document):
    # 主体
    for p in doc.paragraphs:
        yield p
    # 表格
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _clear_paragraph(paragraph):
    # python-docx 没有公开 API 清空段落，只能操作底层 XML
    p_elm = paragraph._p
    for child in list(p_elm):
        p_elm.remove(child)


def _replace_image_in_paragraph(paragraph, placeholder: str, image_path: str, width_cm: Optional[float]) -> bool:
    full = paragraph.text
    if placeholder not in full:
        return False

    # 优先尝试在某个 run 内替换（最不破坏格式）
    for run in paragraph.runs:
        if placeholder in run.text:
            before, after = run.text.split(placeholder, 1)
            run.text = before

            pic_run = paragraph.add_run()
            if width_cm:
                pic_run.add_picture(image_path, width=Cm(width_cm))
            else:
                pic_run.add_picture(image_path)

            if after:
                paragraph.add_run(after)
            return True

    # 兜底：重建段落（可能损失混排，但保证能插入）
    before, _, after = full.partition(placeholder)
    _clear_paragraph(paragraph)
    paragraph.add_run(before)

    pic_run = paragraph.add_run()
    if width_cm:
        pic_run.add_picture(image_path, width=Cm(width_cm))
    else:
        pic_run.add_picture(image_path)

    paragraph.add_run(after)
    return True


def replace_images_in_doc(doc: Document, values_raw: Dict, base_dir: str, image_field_config: Dict[str, Dict]):
    """把形如 {{图片字段}} 的占位符替换成图片。"""
    if not image_field_config:
        return

    for key, cfg in image_field_config.items():
        placeholder = "{{" + key + "}}"
        img_path = _resolve_local_path(base_dir, values_raw.get(key))
        if not img_path or not os.path.exists(img_path):
            continue

        for p in _iter_all_paragraphs(doc):
            if placeholder in p.text:
                _replace_image_in_paragraph(p, placeholder, img_path, cfg.get("width_cm"))


def create_zip_from_folder(folder: str, zip_path: str, arc_root_name: Optional[str] = None) -> None:
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for root, _, files in os.walk(folder):
            for fn in files:
                if fn.startswith("~$"):
                    continue
                full = os.path.join(root, fn)
                rel = os.path.relpath(full, folder)
                arc = os.path.join(arc_root_name, rel) if arc_root_name else rel
                z.write(full, arcname=arc)


def _upload_zip_to_vercel_blob(local_zip_path: str, blob_path: str) -> str:
    """把 zip 上传到 Vercel Blob，返回可下载的 URL。"""
    # 依赖：pip install vercel
    from vercel.blob import BlobClient

    client = BlobClient()  # uses BLOB_READ_WRITE_TOKEN from env
    uploaded = client.upload_file(
        local_zip_path,
        blob_path,
        access="public",
        content_type="application/zip",
    )
    # uploaded.url 一般即可直接下载（zip 会走 attachment）
    return getattr(uploaded, "download_url", None) or getattr(uploaded, "downloadUrl", None) or uploaded.url


# ============ 路由 ============
@app.get("/api/product_types")
def api_product_types():
    if not os.path.isdir(DOCX_TEMPLATE_ROOT):
        return jsonify({"ok": False, "error": "templates 目录不存在"}), 400
    return jsonify({"ok": True, "data": list_product_types()})


@app.get("/api/schema")
def api_schema():
    product_type = request.args.get("product_type", "__root__")
    try:
        folder = get_product_folder(product_type)
    except Exception:
        return jsonify({"ok": False, "error": "product_type 非法"}), 400

    if not os.path.isdir(folder):
        return jsonify({"ok": False, "error": f"未找到产品类型目录：{product_type}"}), 404

    schema = build_schema(product_type)
    return jsonify({"ok": True, "data": schema})


@app.post("/api/generate")
def api_generate():
    payload = request.get_json(force=True, silent=False) or {}

    product_type = payload.get("product_type", "__root__")
    values_raw = payload.get("values", {}) or {}
    mode = payload.get("mode", "all")  # all / selected
    selected = payload.get("selected_templates", []) or []
    blank_unfilled = bool(payload.get("blank_unfilled", True))

    try:
        folder = get_product_folder(product_type)
    except Exception:
        return jsonify({"ok": False, "error": "product_type 非法"}), 400

    tpl_files = list_docx_files(folder)
    if not tpl_files:
        return jsonify({"ok": False, "error": f"该产品类型下没有 .docx 模板：{product_type}"}), 400

    if mode == "selected":
        tpl_files = [f for f in tpl_files if f in selected]
        if not tpl_files:
            return jsonify({"ok": False, "error": "未选择任何模板"}), 400

    schema = build_schema(product_type)
    all_keys = [f["key"] for f in schema.get("fields", [])]

    # ===== 图片字段（可选） =====
    # 例如模板里写 {{指数公司介绍图片}}，前端填写 "static/xxx.png" 或 "presets/xxx.png" 等
    image_field_config: Dict[str, Dict] = {
        # "指数公司介绍图片": {"width_cm": 9.5},
    }
    image_keys = list(image_field_config.keys())

    mapping = build_docx_mapping(
        all_keys=all_keys,
        values_raw=values_raw,
        blank_unfilled=blank_unfilled,
        image_keys=image_keys,
    )

    # run_name：用于 zip 内顶层目录 & blob 路径
    output_folder = (payload.get("output_folder") or "").strip()
    if not output_folder:
        idx = str(values_raw.get("指数名称") or "").strip()
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_folder = f"{product_type}_{idx or '未命名'}_{ts}"

    # Vercel: /tmp 下建立临时目录
    run_dir, run_name = _ensure_unique_subdir(OUTPUT_ROOT, output_folder)

    generated_files: List[str] = []
    zip_path: Optional[str] = None
    try:
        for tpl in tpl_files:
            tpl_path = os.path.join(folder, tpl)
            doc = Document(tpl_path)

            replace_images_in_doc(doc, values_raw=values_raw, base_dir=BASE_DIR, image_field_config=image_field_config)
            replace_placeholders_in_doc(doc, mapping)

            out_name = apply_placeholders_to_filename(tpl, mapping, blank_unfilled=blank_unfilled)
            if not out_name.lower().endswith(".docx"):
                out_name += ".docx"

            out_path = os.path.join(run_dir, out_name)
            doc.save(out_path)
            generated_files.append(out_name)

        # zip
        zip_name = f"{run_name}.zip"
        zip_path = os.path.join(OUTPUT_ROOT, zip_name)
        create_zip_from_folder(run_dir, zip_path, arc_root_name=run_name)

        # 上传到 Blob（推荐：避免 4.5MB 响应体限制 & 无状态文件系统）
        if not os.environ.get("BLOB_READ_WRITE_TOKEN"):
            return (
                jsonify(
                    {
                        "ok": False,
                        "error": "未配置 BLOB_READ_WRITE_TOKEN：Vercel 部署建议使用 Vercel Blob 存储生成的 zip（否则无法可靠下载）。",
                    }
                ),
                500,
            )

        # 为避免覆盖，加入随机后缀
        blob_path = f"generated/{run_name}-{uuid.uuid4().hex}.zip"
        zip_url = _upload_zip_to_vercel_blob(zip_path, blob_path)

        return jsonify(
            {
                "ok": True,
                "data": {
                    "job_id": run_name,
                    "run_name": run_name,
                    "zip_name": zip_name,
                    "zip_url": zip_url,
                    "generated_files": generated_files,
                },
            }
        )
    except Exception as e:
        return jsonify({"ok": False, "error": f"生成失败：{repr(e)}"}), 500
    finally:
        # 清理临时目录（zip 文件也可以删；上传后本地留存无意义）
        try:
            shutil.rmtree(run_dir, ignore_errors=True)
        except Exception:
            pass
        if zip_path:
            try:
                if os.path.exists(zip_path):
                    os.remove(zip_path)
            except Exception:
                pass


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
